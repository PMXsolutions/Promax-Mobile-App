#!/usr/bin/env python3
"""Generate Promax Care Mobile App Project Scope Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date


def set_cell_shading(cell, color_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x03, 0x06, 0x37)
    return heading


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "030637")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Promax Care\nMobile Application")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x03, 0x06, 0x37)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Full Project Scope Document")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Version 1.1.0  |  {date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.italic = True

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run("Powered by Promax IT Solutions\nhttps://promaxsolutions.com.au/")
    run.font.size = Pt(11)

    doc.add_page_break()

    # Table of Contents placeholder
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Project Overview",
        "3. Background and Business Context",
        "4. Project Objectives",
        "5. Scope Definition",
        "6. Stakeholders",
        "7. Functional Requirements",
        "8. Non-Functional Requirements",
        "9. Technical Architecture",
        "10. Technology Stack",
        "11. User Roles and Security",
        "12. Third-Party Integrations",
        "13. Mobile Application Structure",
        "14. API Integration Summary",
        "15. Deployment and DevOps",
        "16. Development Progress",
        "17. Testing Strategy",
        "18. Assumptions and Constraints",
        "19. Risks and Mitigations",
        "20. Success Criteria",
        "21. Future Enhancements",
        "22. Appendices",
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Promax Care is a staff-facing mobile application designed for the Australian disability "
        "and care workforce management sector. Built with React Native and Expo, the application "
        "enables care staff to manage their daily work operations including shift rosters, "
        "attendance tracking with GPS geofencing, shift progress reporting, compliance document "
        "management, availability scheduling, and real-time notifications."
    )
    doc.add_paragraph(
        "This document defines the complete project scope for the Promax Care mobile application "
        "(version 1.1.0), covering functional and non-functional requirements, technical "
        "architecture, integrations, deployment strategy, development progress to date, known gaps, "
        "and success criteria. The mobile client integrates with a separate .NET Core REST API "
        "backend and MySQL database hosted on AWS infrastructure."
    )

    # 2. Project Overview
    add_heading(doc, "2. Project Overview", 1)
    add_table(
        doc,
        ["Attribute", "Detail"],
        [
            ["Project Name", "Promax Care Mobile Application"],
            ["Package Name", "pmax-app"],
            ["Application Version", "1.1.0"],
            ["Platform", "iOS and Android (Expo/React Native)"],
            ["Repository", "https://github.com/pmxsolutions/promax-mobile-app"],
            ["Primary Users", "Care staff (Staff role only)"],
            ["Geographic Context", "Australia (Sydney timezone, NDIS compliance)"],
            ["Backend Platform", ".NET Core REST API with MySQL (separate repository)"],
            ["Deployment", "Expo Application Services (EAS Build, Submit, OTA Updates)"],
        ],
    )

    add_heading(doc, "2.1 Project Purpose", 2)
    doc.add_paragraph(
        "The Promax Care mobile application serves as the primary operational tool for care workers "
        "in the field. It replaces manual paper-based processes and fragmented communication with a "
        "unified digital platform that connects staff to their assigned shifts, clients, compliance "
        "requirements, and reporting obligations in real time."
    )

    add_heading(doc, "2.2 Problem Statement", 2)
    doc.add_paragraph(
        "Care organisations in the Australian NDIS sector face significant operational challenges: "
        "managing distributed workforces across multiple client locations, ensuring accurate "
        "attendance verification, maintaining compliance documentation, and capturing detailed "
        "shift reports for regulatory and billing purposes. Without a dedicated mobile solution, "
        "staff rely on phone calls, paper forms, and web portals that are impractical during active "
        "care delivery."
    )

    add_heading(doc, "2.3 Solution Summary", 2)
    doc.add_paragraph(
        "Promax Care provides a purpose-built mobile experience that allows staff to view their "
        "shift roster on a calendar, clock in and out with GPS-validated location checks, submit "
        "comprehensive shift reports, upload and track compliance documents, manage their weekly "
        "availability, maintain their profile and emergency details, track transport mileage, and "
        "receive push notifications for schedule changes and messages."
    )

    # 3. Background
    add_heading(doc, "3. Background and Business Context", 1)
    doc.add_paragraph(
        "Promax IT Solutions develops workforce management software for the Australian care and "
        "disability services industry. The Promax Care platform comprises a web-based administration "
        "portal, a .NET Core backend API, and this mobile application for field staff."
    )
    add_bullet(doc, " — National Disability Insurance Scheme (NDIS) compliance requirements for worker documentation and reporting.", "NDIS Context")
    add_bullet(doc, " — All shift scheduling, status computation, and clock-in windows use Australia/Sydney timezone.", "Timezone")
    add_bullet(doc, " — Clock-in requires staff to be within approximately 1,000 metres of the client's registered location.", "Geofencing")
    add_bullet(doc, " — Seven predefined compliance document types plus custom uploads (First Aid, Police Check, NDIS Orientation, WWVP Card, Driver's License, Car Insurance, Academic Certificate).", "Document Types")

    # 4. Objectives
    add_heading(doc, "4. Project Objectives", 1)
    objectives = [
        "Enable care staff to view and manage assigned shift rosters from a mobile device.",
        "Provide GPS-validated clock-in and clock-out functionality to ensure attendance accuracy.",
        "Allow staff to submit detailed shift progress reports covering medications, health, goals, incidents, and behaviour.",
        "Support compliance document upload, tracking, and expiration management.",
        "Allow staff to manage weekly availability windows for roster planning.",
        "Provide a complete staff profile management experience including emergency contacts, bank details, and digital signatures.",
        "Deliver real-time push and in-app notifications for schedule changes and organisational messages.",
        "Support transport trip tracking and mileage submission per shift.",
        "Ensure secure JWT-based authentication with role-based access (Staff only on mobile).",
        "Deliver a polished, accessible user experience with offline detection and error recovery.",
    ]
    for obj in objectives:
        add_numbered(doc, obj)

    # 5. Scope
    add_heading(doc, "5. Scope Definition", 1)

    add_heading(doc, "5.1 In Scope", 2)
    in_scope = [
        ("Authentication and Onboarding", "Welcome screens, email/password sign-in, forgot password, OTP email verification, session persistence."),
        ("Shift Management", "Calendar/agenda view, shift detail, live status computation, clock-in/out with geofencing, shift cancellation."),
        ("Shift Reports", "List, search, create, edit, and sign shift progress reports."),
        ("Document Management", "Required compliance checklist, upload/view/edit/delete documents, expiration tracking."),
        ("Staff Profile", "Identity card, personal info, emergency contact, bank info, employment details, social links, digital signature, profile photo."),
        ("Availability Management", "Weekly availability editor with add/edit/delete and overlap validation."),
        ("Notifications", "In-app message inbox, message detail, delete, push notification registration."),
        ("Transport and Mileage", "Trip tracking screen with map, start/end kilometre submission."),
        ("Infrastructure UX", "Offline overlay, pull-to-refresh, error/retry states, haptic feedback."),
        ("Deployment", "EAS Build profiles (development, preview, production), OTA updates, app store submission."),
    ]
    for title, desc in in_scope:
        add_bullet(doc, f" — {desc}", title)

    add_heading(doc, "5.2 Out of Scope", 2)
    out_scope = [
        "Admin web portal functionality (roster creation, client management, invoicing).",
        "Client-facing mobile application.",
        "Backend API development (.NET Core — maintained in separate repository).",
        "Database schema design and MySQL administration.",
        "Xero accounting integration UI (backend-side only).",
        "Invoicing and billing workflows.",
        "Multi-role mobile access (Admin, Client, Care Manager roles use web platform).",
        "End-to-end automated testing (Detox/Maestro).",
        "Continuous Integration pipeline (GitHub Actions).",
    ]
    for item in out_scope:
        add_bullet(doc, item)

    # 6. Stakeholders
    add_heading(doc, "6. Stakeholders", 1)
    add_table(
        doc,
        ["Stakeholder", "Role", "Interest"],
        [
            ["Care Staff", "Primary end users", "Manage shifts, attendance, reports, and documents in the field"],
            ["Care Organisation Admins", "Platform administrators", "Ensure staff compliance and accurate reporting via web portal"],
            ["Clients / Participants", "Service recipients", "Benefit from accurate care delivery and documentation"],
            ["Promax IT Solutions", "Developer / vendor", "Platform delivery, maintenance, and support"],
            ["Regulatory Bodies", "NDIS / compliance", "Accurate documentation and worker qualification records"],
        ],
    )

    # 7. Functional Requirements
    add_heading(doc, "7. Functional Requirements", 1)

    modules = [
        (
            "7.1 Authentication and Onboarding",
            [
                ("FR-AUTH-01", "The app shall display a 3-screen onboarding swiper for first-time users."),
                ("FR-AUTH-02", "Users shall sign in with email and password via POST /Account/auth_login."),
                ("FR-AUTH-03", "Only users with role 'Staff' shall be granted mobile app access; other roles receive a message to use the web platform."),
                ("FR-AUTH-04", "Unconfirmed email accounts shall be redirected to OTP verification (POST /Account/post_otp)."),
                ("FR-AUTH-05", "Users shall request password reset via GET /Account/forgot_password."),
                ("FR-AUTH-06", "JWT tokens shall be persisted securely using expo-secure-store (native) or AsyncStorage (web)."),
                ("FR-AUTH-07", "401 API responses shall trigger automatic logout and cache clearing."),
            ],
        ),
        (
            "7.2 Shift Management",
            [
                ("FR-SHIFT-01", "Staff shall view assigned shifts in a calendar/agenda interface with pull-to-refresh."),
                ("FR-SHIFT-02", "Each shift shall display a computed status: Upcoming, Clock-In, Shift In Progress, Present, Absent, or Cancelled."),
                ("FR-SHIFT-03", "Staff shall view shift details including client profile, activities, description, and map preview."),
                ("FR-SHIFT-04", "Staff shall clock in via GET /Attendances/clock_in with GPS coordinates; clock-in allowed within 10 minutes before shift start and within ~1,000m of client location."),
                ("FR-SHIFT-05", "Staff shall clock out via GET /Attendances/clock_out."),
                ("FR-SHIFT-06", "Staff shall cancel shifts with a reason via GET /ShiftRosters/shift_cancellation."),
                ("FR-SHIFT-07", "A pending-shift banner shall alert staff to completed shifts with unsigned reports."),
            ],
        ),
        (
            "7.3 Shift Reports",
            [
                ("FR-REPORT-01", "Staff shall view a searchable list of shift reports filtered by client name."),
                ("FR-REPORT-02", "Staff shall create shift reports after shift completion covering medication, meal plan, health, goals, incidents, behaviour, and family contact."),
                ("FR-REPORT-03", "Staff shall edit existing shift reports via POST /ShiftReports/edit/{id}."),
                ("FR-REPORT-04", "Reports shall support a signing workflow (isShiftReportSigned flag)."),
            ],
        ),
        (
            "7.4 Document Management",
            [
                ("FR-DOC-01", "Staff shall view a checklist of required compliance documents with status indicators (Pending, Not Submitted, etc.)."),
                ("FR-DOC-02", "Staff shall upload documents via multipart POST /Documents/add_document."),
                ("FR-DOC-03", "Staff shall view, edit, and delete documents."),
                ("FR-DOC-04", "Document expiration dates shall be displayed and tracked."),
                ("FR-DOC-05", "Seven predefined NDIS/compliance document types shall be supported plus custom uploads."),
            ],
        ),
        (
            "7.5 Staff Profile",
            [
                ("FR-PROF-01", "Staff shall view an identity card with profile photo and key details."),
                ("FR-PROF-02", "Staff shall view and edit personal information, emergency contact, and other/social information."),
                ("FR-PROF-03", "Staff shall view bank and employment details (read-only)."),
                ("FR-PROF-04", "Staff shall capture a digital signature via draw, type, or upload methods."),
                ("FR-PROF-05", "Staff shall upload a profile photo."),
                ("FR-PROF-06", "Staff shall log out, clearing session and cached data."),
            ],
        ),
        (
            "7.6 Availability Management",
            [
                ("FR-AVAIL-01", "Staff shall view weekly availability in a day-by-day editor."),
                ("FR-AVAIL-02", "Staff shall add, edit, and delete availability time slots."),
                ("FR-AVAIL-03", "The system shall validate against overlapping time slots."),
            ],
        ),
        (
            "7.7 Notifications",
            [
                ("FR-NOTIF-01", "Staff shall view an in-app message inbox fetched from GET /Messages/inbox."),
                ("FR-NOTIF-02", "Staff shall read message detail content (HTML rendered)."),
                ("FR-NOTIF-03", "Staff shall delete messages."),
                ("FR-NOTIF-04", "The app shall register FCM device tokens with an external push notification service."),
            ],
        ),
        (
            "7.8 Transport and Mileage",
            [
                ("FR-TRANS-01", "Staff shall access a trip tracking screen with map display for active shifts."),
                ("FR-TRANS-02", "Staff shall submit start and end kilometre readings via GET /ShiftRosters/fill_mileage."),
            ],
        ),
    ]

    for section_title, requirements in modules:
        add_heading(doc, section_title, 2)
        add_table(doc, ["ID", "Requirement"], requirements)

    # 8. Non-Functional Requirements
    add_heading(doc, "8. Non-Functional Requirements", 1)
    add_table(
        doc,
        ["ID", "Category", "Requirement"],
        [
            ["NFR-01", "Performance", "Lists shall use FlashList for optimised rendering of large datasets."],
            ["NFR-02", "Performance", "Server state shall be cached via TanStack React Query with stale-while-revalidate."],
            ["NFR-03", "Security", "Authentication tokens stored in expo-secure-store on native platforms."],
            ["NFR-04", "Security", "All API communication over HTTPS with Bearer JWT authorization."],
            ["NFR-05", "Reliability", "Network connectivity monitored; offline overlay displayed when disconnected."],
            ["NFR-06", "Usability", "Forms validated with Zod schemas via React Hook Form."],
            ["NFR-07", "Usability", "Toast notifications via react-native-flash-message for user feedback."],
            ["NFR-08", "Accessibility", "Haptic feedback on tab navigation (iOS)."],
            ["NFR-09", "Maintainability", "TypeScript strict mode with typed Expo Router routes."],
            ["NFR-10", "Deployability", "OTA updates via Expo Updates for production channel."],
            ["NFR-11", "Compatibility", "iOS (phone only, no tablet) and Android support via Expo SDK 53."],
            ["NFR-12", "Localisation", "Business logic uses Australia/Sydney timezone throughout."],
        ],
    )

    # 9. Technical Architecture
    add_heading(doc, "9. Technical Architecture", 1)
    doc.add_paragraph(
        "The Promax Care mobile application follows a client-server architecture. The mobile client "
        "is a React Native application built with Expo, communicating with a remote .NET Core REST "
        "API. State management is split between Zustand (client/auth state with persistence) and "
        "TanStack React Query (server state with caching). Navigation uses Expo Router file-based "
        "routing with typed routes."
    )

    add_heading(doc, "9.1 Architecture Layers", 2)
    add_table(
        doc,
        ["Layer", "Components", "Responsibility"],
        [
            ["Presentation", "app/, components/, modules/", "Screens, UI components, form modules"],
            ["State Management", "store/, hooks/", "Zustand stores, React Query hooks"],
            ["Service Layer", "services/, libs/", "API calls, Axios instance, interceptors"],
            ["Domain", "types/, helpers/, constants/", "TypeScript models, business logic, theme"],
            ["Infrastructure", "context/, utils/", "Network provider, storage, file handling, push helpers"],
        ],
    )

    add_heading(doc, "9.2 Data Flow", 2)
    flow_steps = [
        "User interacts with a screen component in the app/ directory.",
        "Screen invokes a React Query hook (hooks/) which calls a service function (services/).",
        "Service function uses the Axios instance (libs/) with JWT interceptor to call the REST API.",
        "Response data is cached by React Query and rendered in the UI.",
        "Mutations invalidate relevant query keys to refresh dependent data.",
        "Auth state changes propagate via Zustand store, triggering route guards in auth-wrapper.",
    ]
    for step in flow_steps:
        add_numbered(doc, step)

    # 10. Technology Stack
    add_heading(doc, "10. Technology Stack", 1)

    add_heading(doc, "10.1 Mobile Frontend", 2)
    add_table(
        doc,
        ["Category", "Technology", "Version"],
        [
            ["Framework", "React Native", "0.79.5"],
            ["UI Library", "React", "19.0.0"],
            ["Platform", "Expo SDK", "53.0.20"],
            ["Routing", "Expo Router", "5.1.4"],
            ["Language", "TypeScript", "5.3.3"],
            ["State (Client)", "Zustand", "5.0.2"],
            ["State (Server)", "TanStack React Query", "5.81.5"],
            ["HTTP Client", "Axios", "1.7.9"],
            ["Forms", "React Hook Form + Zod", "7.54.0 / 3.24.0"],
            ["Maps", "react-native-maps, Google Places", "1.20.1"],
            ["Location", "expo-location, geolib", "18.1.6 / 3.3.4"],
            ["Storage", "expo-secure-store, AsyncStorage", "14.2.3 / 2.1.2"],
            ["Notifications", "expo-notifications", "0.31.4"],
            ["UI Components", "FlashList, Bottom Sheet, Lottie, Reanimated", "Various"],
            ["Dates", "date-fns, date-fns-tz", "4.1.0 / 3.2.0"],
            ["Testing", "Jest + jest-expo", "29.2.1"],
        ],
    )

    add_heading(doc, "10.2 Backend (External)", 2)
    add_table(
        doc,
        ["Category", "Technology"],
        [
            ["API Framework", ".NET Core REST API"],
            ["Database", "MySQL"],
            ["Authentication", "JSON Web Tokens (JWT)"],
            ["Hosting", "AWS with Docker containerisation"],
            ["Default API Host", "profitmax-001-site10.ctempurl.com/api"],
        ],
    )

    # 11. User Roles and Security
    add_heading(doc, "11. User Roles and Security", 1)

    add_heading(doc, "11.1 Role-Based Access", 2)
    add_table(
        doc,
        ["Role", "Mobile Access", "Platform"],
        [
            ["Staff", "Full application access", "Mobile app (this project)"],
            ["Admin", "Blocked — directed to web", "Web portal"],
            ["Client", "Blocked — directed to web", "Web portal"],
            ["Other roles", "Blocked — directed to web", "Web portal"],
        ],
    )

    add_heading(doc, "11.2 Authentication Flow", 2)
    auth_steps = [
        "User submits email/password on sign-in screen.",
        "App sends POST /Account/auth_login and receives JWT token, userProfile, and staffProfile.",
        "Role validation: only 'Staff' role proceeds; others see web platform message.",
        "Token and profile stored in Zustand, persisted to SecureStore.",
        "Axios interceptor attaches Authorization: Bearer header to all subsequent requests.",
        "On 401 response, automatic logout clears token and React Query cache.",
    ]
    for step in auth_steps:
        add_numbered(doc, step)

    add_heading(doc, "11.3 Permissions Required", 2)
    add_table(
        doc,
        ["Permission", "Platform", "Purpose"],
        [
            ["Camera", "iOS", "Profile photo and document capture"],
            ["Photo Library", "iOS", "Image upload from gallery"],
            ["Location (When In Use)", "iOS / Android", "Clock-in geofencing, trip tracking, maps"],
            ["Fine/Coarse Location", "Android", "GPS attendance validation"],
            ["Notifications", "iOS / Android", "Push notification delivery"],
        ],
    )

    # 12. Integrations
    add_heading(doc, "12. Third-Party Integrations", 1)
    add_table(
        doc,
        ["Integration", "Purpose", "Endpoint / Service"],
        [
            ["Promax REST API", "All business data", "EXPO_PUBLIC_API_BASEURL"],
            ["Google Maps Platform", "Maps, Places autocomplete, directions, static maps", "EXPO_PUBLIC_GOOGLE_MAPS_API_KEY"],
            ["Firebase Cloud Messaging", "Android push notifications", "google-services.json"],
            ["Expo Push Notifications", "Token registration, local notifications", "Expo SDK"],
            ["Render.com Push Service", "FCM token storage per user/company", "push-notification-r5mb.onrender.com"],
            ["Expo Application Services", "Builds, OTA updates, store submission", "EAS project ID: 70d0de46-..."],
            ["Xero (Backend)", "Accounting integration", "Backend-side only (xeroUploadLink)"],
        ],
    )

    # 13. App Structure
    add_heading(doc, "13. Mobile Application Structure", 1)

    add_heading(doc, "13.1 Navigation Architecture", 2)
    doc.add_paragraph(
        "The app uses Expo Router with two primary route groups: (auth) for unauthenticated flows "
        "and (root) for the authenticated application. The authenticated area includes a bottom tab "
        "navigator with five tabs and a stack navigator for detail screens."
    )

    add_heading(doc, "13.2 Screen Inventory", 2)
    add_table(
        doc,
        ["Route Group", "Screen", "Description"],
        [
            ["Auth", "/(auth)/welcome", "3-screen onboarding swiper"],
            ["Auth", "/(auth)/sign-in", "Email/password login"],
            ["Auth", "/(auth)/forgot-password", "Password reset request"],
            ["Auth", "/(auth)/otp-verification/[email]", "Email OTP verification"],
            ["Auth", "/(auth)/change-password/[email]", "Password change form"],
            ["Tabs", "/(root)/(tabs)/index", "Shift roster calendar"],
            ["Tabs", "/(root)/(tabs)/reports", "Shift reports list"],
            ["Tabs", "/(root)/(tabs)/add", "Weekly availability editor"],
            ["Tabs", "/(root)/(tabs)/documents", "Compliance documents"],
            ["Tabs", "/(root)/(tabs)/profile", "Staff profile hub"],
            ["Stack", "/(root)/shift", "Shift detail view"],
            ["Stack", "/(root)/shift/pending", "Pending unsigned reports"],
            ["Stack", "/(root)/shift/cancel", "Shift cancellation form"],
            ["Stack", "/(root)/report/create", "Create/edit shift report"],
            ["Stack", "/(root)/document/[id]", "Document detail"],
            ["Stack", "/(root)/document/add-document", "Upload new document"],
            ["Stack", "/(root)/profile/*", "Profile sub-screens (6 sections + edit forms)"],
            ["Stack", "/(root)/notification", "Message inbox and detail"],
            ["Stack", "/(root)/transport/trip-screen", "Trip tracking with map"],
        ],
    )

    add_heading(doc, "13.3 Directory Structure", 2)
    dirs = [
        ("app/", "Expo Router screens and layouts"),
        ("components/", "Reusable UI (shared, shift, profile, map, transport, signature)"),
        ("modules/", "Feature form modules (auth, profile, shift, report, document)"),
        ("services/", "API service layer"),
        ("hooks/", "React Query hooks and custom hooks"),
        ("store/", "Zustand stores (auth, location)"),
        ("types/", "TypeScript domain models"),
        ("helpers/", "Shift status logic and actions"),
        ("constants/", "Theme, storage keys, onboarding data"),
        ("utils/", "Storage, file handling, push helpers"),
        ("libs/", "Axios instance, React Query client"),
        ("context/", "NetworkProvider"),
        ("assets/", "Images and Lottie animations"),
        ("__tests__/", "Jest unit tests"),
    ]
    add_table(doc, ["Directory", "Purpose"], dirs)

    # 14. API Integration
    add_heading(doc, "14. API Integration Summary", 1)
    doc.add_paragraph(
        "The mobile application consumes REST endpoints from the Promax backend API. "
        "All authenticated requests include a Bearer JWT token. The base URL is configured "
        "via the EXPO_PUBLIC_API_BASEURL environment variable."
    )

    api_groups = [
        ("Account", "auth_login, forgot_password, resend_otp, post_otp"),
        ("Shift Rosters", "get_shifts_by_user, get by ID, shift_cancellation, fill_mileage"),
        ("Attendances", "clock_in, clock_out"),
        ("Shift Reports", "get_staff_shiftreports, get_shiftreport_details, add_shiftreport, edit"),
        ("Documents", "get_all_staff_documents, get_document, add_document, edit, delete"),
        ("Staffs", "get by ID, edit (multipart)"),
        ("Staff Availabilities", "get_staff_availabilities, add, edit, delete"),
        ("Companies", "get_company"),
        ("Messages", "inbox, get by ID, delete"),
    ]
    add_table(doc, ["API Group", "Endpoints"], api_groups)

    # 15. Deployment
    add_heading(doc, "15. Deployment and DevOps", 1)

    add_heading(doc, "15.1 EAS Build Profiles", 2)
    add_table(
        doc,
        ["Profile", "Purpose", "Configuration"],
        [
            ["development", "Local development with dev client", "iOS simulator, internal distribution"],
            ["preview", "QA and stakeholder testing", "Android APK, preview channel"],
            ["production", "App store release", "Android app bundle, iOS auto-increment, production channel"],
        ],
    )

    add_heading(doc, "15.2 Environment Configuration", 2)
    add_table(
        doc,
        ["Variable", "Purpose", "Required"],
        [
            ["EXPO_PUBLIC_API_BASEURL", "Backend REST API base URL", "Yes"],
            ["EXPO_PUBLIC_GOOGLE_MAPS_API_KEY", "Google Maps and Places API", "Yes"],
            ["google-services.json", "Firebase/FCM configuration (Android)", "Yes (not in repo)"],
        ],
    )

    add_heading(doc, "15.3 App Identifiers", 2)
    add_table(
        doc,
        ["Platform", "Identifier"],
        [
            ["iOS Bundle ID", "com.promax-app"],
            ["Android Package", "com.tobby95.pmaxapp"],
            ["Expo Slug", "promax-care"],
            ["EAS Project ID", "70d0de46-7142-4d9a-a7b5-a089e9ea2879"],
            ["OTA Runtime Version", "1.0.0"],
        ],
    )

    add_heading(doc, "15.4 Build Commands", 2)
    commands = [
        "npm start — Start Expo development server",
        "npm run android — Run on Android device/emulator",
        "npm run ios — Run on iOS simulator",
        "npm run web — Start web development server",
        "npm test — Run Jest tests",
        "npm run lint — Run Expo linter",
        "eas build --profile production — Production build via EAS",
    ]
    for cmd in commands:
        add_bullet(doc, cmd)

    # 16. Development Progress
    add_heading(doc, "16. Development Progress", 1)
    doc.add_paragraph(
        "The application has progressed through iterative development cycles covering all major "
        "feature modules. The following milestones represent the development trajectory:"
    )

    milestones = [
        ("Phase 1 — Foundation", "Project scaffolding, Expo Router setup, authentication flow, Zustand auth store, Axios configuration."),
        ("Phase 2 — Shift Management", "Shift roster calendar, shift detail views, clock-in/out with GPS geofencing, shift status computation, cancellation flow."),
        ("Phase 3 — Reporting", "Shift report list, create/edit forms with comprehensive fields, report signing workflow, pending report alerts."),
        ("Phase 4 — Documents", "Compliance document checklist, upload/view/edit/delete, predefined NDIS document types, expiration tracking."),
        ("Phase 5 — Profile", "Identity card, personal/emergency/bank/employment info screens, digital signature capture (draw/type/upload), profile photo upload."),
        ("Phase 6 — Availability", "Weekly availability editor with day-by-day time slots, overlap validation, CRUD operations."),
        ("Phase 7 — Notifications", "In-app message inbox, push notification registration via FCM and external token service."),
        ("Phase 8 — Transport", "Trip tracking screen with map, mileage submission."),
        ("Phase 9 — Polish", "Network connectivity overlay, Lottie animations, font updates (Inter), onboarding screens, haptic feedback, error/retry states."),
        ("Phase 10 — Security Hardening", "Auth fixes, attendance validation, report integrity, secret exposure remediation (PR #2)."),
    ]
    add_table(doc, ["Phase", "Deliverables"], milestones)

    add_heading(doc, "16.1 Recent Changes", 2)
    recent = [
        "Fixed mobile auth, attendance, report, and secret exposure issues (PR #2).",
        "Report screen improvements and bug fixes.",
        "iOS deployment error resolution.",
        "Google Places autocomplete input fixes.",
        "App rebranding to Promax Care.",
        "Shift calendar UI enhancements.",
        "Digital signature feature addition.",
        "Network connectivity tab/overlay.",
        "Geofencing distance adjustments.",
    ]
    for item in recent:
        add_bullet(doc, item)

    # 17. Testing
    add_heading(doc, "17. Testing Strategy", 1)

    add_heading(doc, "17.1 Current Test Coverage", 2)
    add_table(
        doc,
        ["Test File", "Scope", "Framework"],
        [
            ["__tests__/shift-service.test.ts", "getActivityDetailStatus() — Absent, Clock-In window, Cancelled shifts", "Jest + jest-expo"],
        ],
    )

    add_heading(doc, "17.2 Testing Gaps", 2)
    gaps = [
        "No integration tests for API service layer.",
        "No end-to-end tests (Detox, Maestro, or similar).",
        "No component/snapshot tests.",
        "No CI pipeline to run tests automatically.",
        "Limited unit test coverage (shift status logic only).",
    ]
    for gap in gaps:
        add_bullet(doc, gap)

    add_heading(doc, "17.3 Recommended Testing Roadmap", 2)
    roadmap = [
        "Expand unit tests for helpers/, services/, and store/ modules.",
        "Add React Query hook tests with mock service layer.",
        "Implement component tests for critical forms (sign-in, clock-in, report creation).",
        "Set up GitHub Actions CI pipeline for lint and test on pull requests.",
        "Evaluate Maestro or Detox for critical user flow E2E testing.",
    ]
    for item in roadmap:
        add_numbered(doc, item)

    # 18. Assumptions
    add_heading(doc, "18. Assumptions and Constraints", 1)

    add_heading(doc, "18.1 Assumptions", 2)
    assumptions = [
        "The .NET Core backend API is available, stable, and maintained in a separate repository.",
        "Staff users have smartphones with GPS capability and internet connectivity during shifts.",
        "Google Maps API key and Firebase configuration are provisioned for production.",
        "EAS credentials and app store accounts are managed by Promax IT Solutions.",
        "Backend enforces server-side authorization; mobile client performs role check at login only.",
        "Care organisations operate within Australia/Sydney timezone.",
    ]
    for item in assumptions:
        add_bullet(doc, item)

    add_heading(doc, "18.2 Constraints", 2)
    constraints = [
        "Mobile app restricted to Staff role only — no admin or client functionality.",
        "Clock-in geofencing limited to ~1,000 metres from client registered location.",
        "iOS deployment limited to iPhone (no iPad support).",
        "No offline-first data sync — app requires network connectivity for all operations.",
        "Change password API integration not yet wired (UI exists, backend call stubbed).",
        "Query client persistence prepared but currently commented out.",
    ]
    for item in constraints:
        add_bullet(doc, item)

    # 19. Risks
    add_heading(doc, "19. Risks and Mitigations", 1)
    add_table(
        doc,
        ["Risk", "Impact", "Likelihood", "Mitigation"],
        [
            ["Backend API downtime", "High — app unusable", "Medium", "Offline overlay alerts users; retry mechanisms on API calls"],
            ["GPS inaccuracy affecting clock-in", "High — staff cannot clock in", "Medium", "1,000m geofence radius provides tolerance; manual override via web portal"],
            ["Minimal test coverage", "Medium — regressions undetected", "High", "Expand unit tests; implement CI pipeline"],
            ["Secret/credential exposure", "Critical — security breach", "Low", "Addressed in PR #2; env vars gitignored; no secrets in repo"],
            ["Google Maps API quota/cost", "Medium — maps unavailable", "Low", "Monitor usage; implement static map fallbacks"],
            ["App store rejection", "Medium — delayed release", "Low", "Follow platform guidelines; EAS submit automation"],
            ["Incomplete change password flow", "Low — users cannot reset via app", "Medium", "Wire API endpoint; test full recovery flow"],
            ["Single timezone assumption", "Low — incorrect shift times", "Low", "Document Sydney timezone requirement; validate with backend"],
        ],
    )

    # 20. Success Criteria
    add_heading(doc, "20. Success Criteria", 1)
    criteria = [
        "Staff can sign in, view shifts, and clock in/out with GPS validation successfully.",
        "Shift reports can be created, edited, and signed within the mobile app.",
        "All seven compliance document types can be uploaded and tracked.",
        "Staff availability can be managed with valid time slot CRUD operations.",
        "Push notifications are delivered and in-app messages are accessible.",
        "Application builds and deploys successfully via EAS for both iOS and Android.",
        "OTA updates deliver to production channel without requiring app store resubmission.",
        "No critical or high-severity security vulnerabilities in authentication or data handling.",
        "Application achieves stable performance with no crashes during core user flows.",
    ]
    for item in criteria:
        add_numbered(doc, item)

    # 21. Future Enhancements
    add_heading(doc, "21. Future Enhancements", 1)
    enhancements = [
        ("Change Password API Integration", "Wire the change-password form to the backend API endpoint."),
        ("Query Client Persistence", "Enable TanStack React Query persistence for offline cache recovery."),
        ("Expanded Test Coverage", "Unit, integration, and E2E tests with CI pipeline."),
        ("Offline-First Sync", "Queue mutations when offline and sync when connectivity restored."),
        ("Biometric Authentication", "Face ID / fingerprint login for returning users."),
        ("Multi-Language Support", "Internationalisation for diverse care workforce."),
        ("Dark Mode", "Full dark theme support (userInterfaceStyle: automatic configured)."),
        ("In-App Chat", "Real-time messaging between staff and care managers."),
        ("Shift Swap Requests", "Allow staff to request shift swaps from mobile."),
        ("Analytics Dashboard", "In-app performance metrics for staff (hours worked, reports submitted)."),
    ]
    add_table(doc, ["Enhancement", "Description"], enhancements)

    # 22. Appendices
    add_heading(doc, "22. Appendices", 1)

    add_heading(doc, "Appendix A — Domain Entity Summary", 2)
    add_table(
        doc,
        ["Entity", "Key Fields"],
        [
            ["UserProfile", "userId, email, role, companyId, token, tokenExpiration"],
            ["StaffProfile", "Personal, employment, bank, emergency, social, signature, isAdmin, auditApproved"],
            ["ShiftRoster", "shiftRosterId, dates, status, attendance, isEnded, client profile, activities"],
            ["ClientProfile", "NDIS fields, address, lat/lng, care manager, next of kin"],
            ["ShiftReport", "Medication, meal plan, health, goals, incidents, behaviour fields"],
            ["Document", "documentName, status, expirationDate, file URL, userRole"],
            ["StaffAvailability", "Day, fromTimeOfDay, toTimeOfDay"],
            ["Notification/Message", "subject, content, emailFrom/To, attachments"],
            ["Company", "ABN, subscription package, emails, signature"],
        ],
    )

    add_heading(doc, "Appendix B — Predefined Compliance Documents", 2)
    docs_list = [
        "First Aid Certificate",
        "Police Check",
        "NDIS Orientation Certificate",
        "Working with Vulnerable People (WWVP) Card",
        "Driver's License",
        "Car Insurance",
        "Academic Certificate",
    ]
    for d in docs_list:
        add_bullet(doc, d)

    add_heading(doc, "Appendix C — Shift Status Definitions", 2)
    add_table(
        doc,
        ["Status", "Description"],
        [
            ["Upcoming", "Shift scheduled but clock-in window not yet open"],
            ["Clock-In", "Within 10-minute pre-shift window; staff can clock in"],
            ["Shift In Progress", "Staff has clocked in; shift is active"],
            ["Present", "Staff attended and shift completed normally"],
            ["Absent", "Staff did not clock in and shift time has passed"],
            ["Cancelled", "Shift was cancelled with a reason"],
        ],
    )

    add_heading(doc, "Appendix D — Document Control", 2)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Document Title", "Promax Care Mobile Application — Project Scope"],
            ["Version", "1.0"],
            ["Date", date.today().strftime("%B %d, %Y")],
            ["Application Version", "1.1.0"],
            ["Author", "Promax IT Solutions"],
            ["Status", "Draft"],
        ],
    )

    return doc


if __name__ == "__main__":
    import os

    output_dir = os.path.join(os.path.dirname(__file__))
    output_path = os.path.join(output_dir, "Promax_Care_Project_Scope.docx")

    document = build_document()
    document.save(output_path)
    print(f"Document saved to: {output_path}")
